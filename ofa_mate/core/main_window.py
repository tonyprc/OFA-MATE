#!/usr/bin/python3
# -*- coding: utf-8 -*-
# OFA Mate V.1.0.0 for OFA ParaConc
# Copyright (c) 2020 Tony Chang (42716403@qq.com)

# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.

# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.

# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

import os, re,json, xlrd
from docx import Document
from collections import defaultdict
from PyQt5.QtWidgets import QFileDialog

from ofa_mate.ui.ui_main_window import UIMainWindow
from ofa_mate.core.info_collector import InfoCollector
from ofa_mate.core.chapter_filler import ChapterFiller
from ofa_mate.core.list_preparer import ListPreparer
from ofa_mate.core.lang_detector import LangDetector

class MainWindow:
    '''Play the role of a coordinator'''
    def __init__(self):
        currentDir = os.getcwd()
        dataDir = os.path.join(currentDir, "app_data")
        workFileDir = os.path.join(dataDir, "workfiles")
        self._outPutDir = os.path.join(currentDir, "savedfiles")
        self._interface_lang_file = os.path.join(workFileDir,'interface_language_setting.txt')
        self._interface_lang_dict = os.path.join(workFileDir,'interface_language_dict.json')
        self.fc_lg, self.fc_dict = self.set_lang()


        # recieve signals
        self._ui = UIMainWindow()
        self._ui.open_file.connect(self.open_file)
        self._ui.reload_file.connect(self.reload_file)        
        self._ui.load_next_version.connect(self.load_next_version)
        self._ui.write_to_json.connect(self.write_to_json)
        self._ui.json_sl_filler.connect(self.json_sl_submitter)
        self._ui.json_tl_filler.connect(self.json_tl_submitter)
        self._ui.json_sl_refiller.connect(self.json_sl_refiller)
        self._ui.json_tl_refiller.connect(self.json_tl_refiller)

        self._info_collector = InfoCollector()
        self._list_preparer = ListPreparer()
        self._lg_detect = LangDetector()
        self._chapter_maker = ChapterFiller()

        # data preparation
        self.sl = "en"
        self.tl = 'zh'
        self._opt_dict = {}
        self._temp_list = []
        self._temp_tl_list = []
        self._file_suffix = ""

        self._current_file = ''
        self._current_file_list = []
        self._file_id_list = []
        
        self._current_sl_para_list = []
        self._current_tl_para_list = []
        self._current_sl_chapter_num_list = []

        self._sl_bookDict = {}
        self._tl_bookDict_list = []

        self._current_dict_key = ""
        self._current_sl_dict_version = ""
        self._current_tl_dict_version = "" 

        self._ui.show()

    def set_lang(self):
        with open (self._interface_lang_file, mode = 'r', encoding = 'utf-8-sig') as f:
            default_lg = f.read().strip()
        with open (self._interface_lang_dict, mode = 'r', encoding = 'utf-8-sig') as f:
            lg_dict = json.loads(f.read())
        return default_lg, lg_dict

    def open_single_file(self, target_file):
        if self._ui._file_openBox.text():
            self.form_reset()
            if target_file.endswith('txt'):
                with open(target_file, 'rt', encoding = 'utf-8-sig') as f:
                    text = f.read()
                    if "<seg" in text:
                        self._ui._file_cuc_mark_box.setChecked(True)
                        self._opt_dict['marker_seg'] = 1
                        text = re.sub(r'<seg.*"(\d+)"?>(.*)?</seg>', '\g<1>\t\g<2>', text)
                    else:
                        self._ui._file_cuc_mark_box.setChecked(False)
                        self._opt_dict['marker_seg'] = 0
                    para_list = text.split('\n')
                if para_list:
                    self.sg_file_evaluator(para_list)
                else:
                    pass
            elif target_file.endswith('docx'):
                file_pos = self._opt_dict['lang_pos']
                form_marker = self._opt_dict['marker_table']
                self._ui._file_cuc_mark_box.setChecked(False)
                self._opt_dict['marker_seg'] = 0
                self._ui._file_chapt_num_box.setChecked(False)
                self._opt_dict['marker_chapt'] = 0
                self._ui._file_portion_box.setValue(1)
                self._opt_dict['lang_cols'] = 1

                # 无表格标记
                if form_marker == 0:
                    self._ui._file_pos_box.setCurrentIndex(0)
                    self._opt_dict['lang_pos'] = self.fc_dict['u_d'][self.fc_lg]
                    self._ui._file_tab_mark_box.setChecked(False)
                    self._opt_dict['marker_tab'] = 0
                    with open(target_file, 'rb') as doc:
                        document = Document(doc)
                        para_list = [para.text.strip() for para in document.paragraphs]
                    # 列表为空字串，转为靠长度判断其是否为有效列表
                    list_length = len(para_list)
                    if list_length > 1:
                        pass
                    else:
                        # 列表为空字串，自动补填表格标记并更新字典值，直接尝试读取表格
                        para_list = []
                        with open(target_file, 'rb') as doc:
                            document = Document(doc)
                            tables = document.tables
                            for table in tables[:]:
                                for i, row in enumerate(table.rows[:]):
                                    row_content = []
                                    for cell in row.cells[:]:
                                        c = cell.text
                                        row_content.append(c)
                                    row_string = '\t'.join(row_content)
                                    para_list.append(row_string)
                        list_length = len(para_list)
                        if list_length > 1:
                            self._ui._file_table_mark_box.setChecked(True)
                            self._opt_dict['marker_table'] = 1
                            self._ui._file_pos_box.setCurrentIndex(2)
                            self._opt_dict['lang_pos'] =self.fc_dict['l_r'][self.fc_lg]
                            self._ui._file_tab_mark_box.setChecked(False)
                            self._opt_dict['marker_tab'] = 0
                        else:
                            para_list = []
                # 有表格标记
                else:
                    self._ui._file_pos_box.setCurrentIndex(2)
                    self._opt_dict['lang_pos'] = self.fc_dict['l_r'][self.fc_lg]
                    self._ui._file_tab_mark_box.setChecked(False)
                    self._opt_dict['marker_tab'] = 0
                    para_list = []
                    with open(target_file, 'rb') as doc:
                        document = Document(doc)
                        tables = document.tables
                        for table in tables[:]:
                            for i, row in enumerate(table.rows[:]):
                                row_content = []
                                for cell in row.cells[:]:
                                    c = cell.text
                                    row_content.append(c)
                                row_string = '\t'.join(row_content)
                                para_list.append(row_string)
                    list_length = len(para_list)
                    if list_length > 1:
                        pass
                    else:
                        with open(target_file, 'rb') as doc:
                            document = Document(doc)
                            para_list = [para.text.strip() for para in document.paragraphs]
                        list_length = len(para_list)
                        if list_length > 1:
                            self._ui._file_pos_box.setCurrentIndex(0)
                            self._opt_dict['lang_pos'] = self.fc_dict['u_d'][self.fc_lg]
                            self._ui._file_tab_mark_box.setChecked(False)
                            self._opt_dict['marker_tab'] = 0
                            self._ui._file_table_mark_box.setChecked(False)
                            self._opt_dict['marker_table'] = 0
                        else:
                            para_list = []
                if para_list:
                    self.sg_file_evaluator(para_list)
                else:
                    self._ui._set_status_text(self.fc_dict["warning_read_fail_unknown"][self.fc_lg])
            elif target_file.endswith('xlsx'):
                self._ui._file_pos_box.setCurrentIndex(1)
                self._opt_dict['lang_pos'] = self.fc_dict['l_r'][self.fc_lg]
                self._ui._file_table_mark_box.setChecked(True)
                self._opt_dict['marker_table'] = 1
                loc = (target_file)
                wb = xlrd.open_workbook(loc)
                sheet_id = wb.sheet_names()
                temp_para_list = []
                for v, text_id in enumerate(sheet_id):
                    sheet = wb.sheet_by_name(text_id)
                    content = []
                    for i in range(sheet.nrows):
                        content_line = []
                        for j in range(sheet.ncols):
                            line_content = sheet.cell_value(i, j)
                            # 防止序号为浮点数时显示小数点位
                            if isinstance(line_content, float):
                                line_content = str(int(line_content))
                            else:
                                pass
                            content_line.append(line_content)
                        content.append("\t".join(content_line))
                    temp_para_list.append(content)
                if temp_para_list:
                    for para_list in temp_para_list:
                        self.sg_file_evaluator(para_list)
                elif self._current_sl_para_list:
                    para_list = self._current_sl_para_list
                    self.sg_file_evaluator(para_list)
                else:
                    self._ui._set_status_text(self.fc_dict["warning_read_fail_unknown"][self.fc_lg])
            else:
                pass
        else:
            pass

    # file_open_group
    def open_multi_file(self, target_files, file_id_list):
        if self._ui._file_openBox.text():
            self.form_reset()
            if target_files:
                file_groups = []
                for filename in set(file_id_list):
                    s_list = []
                    for filepath in target_files:
                        if filename in filepath:
                            s_list.append(filepath)
                        else:
                            pass
                    file_groups.append(s_list)
                for file_group in file_groups:
                    file_num = len(file_group)
                    self._ui._file_num_box.setValue(file_num)
                    self.form_reset()
                    sl_para_list = []
                    tl_para_list = []
                    for file in file_group:
                        if 's0' in file:
                            with open(file, 'rt', encoding = 'utf-8-sig') as f:
                                text = f.read()
                                if "<seg" in text:
                                    self._ui._file_cuc_mark_box.setChecked(True)
                                    self._opt_dict['marker_seg'] = 1
                                    text = re.sub(r'<seg.*"(\d+)"?>(.*)?</seg>', '\g<1>\t\g<2>', text)
                                else:
                                    self._opt_dict['marker_seg'] = 0
                            para_list = text.split('\n')
                            sl_para_list.extend([sent.strip().replace('ZZZZZ.', "") for sent in para_list])
                        else:
                            with open(file, 'rt', encoding = 'utf-8-sig') as f:
                                text = f.read()
                                if "<seg" in text:
                                    self._ui._file_cuc_mark_box.setChecked(True)
                                    self._opt_dict['marker_seg'] = 1
                                    text = re.sub(r'<seg.*"(\d+)"?>(.*)?</seg>', '\g<1>\t\g<2>', text)
                                else:
                                    self._opt_dict['marker_seg'] = 0
                            para_list = text.split('\n')
                            tl_para_list.append([sent.strip() for sent in para_list])
                    self._current_sl_para_list = sl_para_list
                    self._current_tl_para_list = tl_para_list
                    if sl_para_list:
                        self.sep_files_evaluator(sl_para_list)
                        self._current_sl_para_list, self._current_tl_para_list = self._list_preparer.prepare_seperate_bi_list(self._ui, self._opt_dict, self.sl, self.tl, sl_para_list, tl_para_list)
                    else:
                        self._current_file_list.clear()
                        self._ui._set_status_text(self.fc_dict["warning_naming_error_mul"][self.fc_lg])
            else:
                pass
        else:
            pass
        return self._current_sl_para_list, self._current_tl_para_list

    # file_open_group
    def reload_file(self):
        if self._ui._file_openBox.text():
            current_sl_text = self._ui._ss_book_contentsBox.toPlainText()
            if current_sl_text:
                self._ui._set_status_text(self.fc_dict["warning_overload"][self.fc_lg])
            else:
                self._opt_dict = self.opt_checker()
                file_type = self._opt_dict['file_type']
                if self._ui._file_openButton.text() == self.fc_dict["open_s"][self.fc_lg]:
                    if self._current_file:
                        self.open_single_file(self._current_file)
                elif self._ui._file_openButton.text() == self.fc_dict["open_m"][self.fc_lg]:
                    if self._current_file_list and self._file_id_list:
                        self.open_multi_file(self._current_file_list, self._file_id_list)
                    else:
                        self._current_file_list.clear()
                        self._file_id_list.clear()
                        self._ui._set_status_text(self.fc_dict["warning_naming_error"][self.fc_lg])
                else:
                    pass
        else:
            self._ui._set_status_text(self.fc_dict["warning_no_file"][self.fc_lg])

    # 打开单文件，只读取目标文件夹名称及路径，同时生成操作选项核对列表
    # file_open_group
    def open_file(self):
        self._opt_dict.clear()
        self._opt_dict = self.opt_checker()
        self._file_suffix = self._opt_dict['file_type']
        self.form_reset()
        file_type = self._opt_dict['file_type']
        if self._ui._file_openButton.text() == self.fc_dict["open_s"][self.fc_lg]:
            options = QFileDialog.Options()
            options |= QFileDialog.DontUseNativeDialog
            presentFile = QFileDialog.getOpenFileName(self._ui, self.fc_dict["open_s"][self.fc_lg], "", f"*.{self._file_suffix}", options = options)
            if presentFile:
                present_file_path = presentFile[0]
                present_file = present_file_path.split('/')[-1]
                self._ui._file_openBox.setText(present_file)
                self._current_file = present_file_path
                self.open_single_file(self._current_file)
        elif self._ui._file_openButton.text() == self.fc_dict["open_m"][self.fc_lg]:
            options = QFileDialog.Options()
            options |= QFileDialog.ShowDirsOnly
            presentDir = QFileDialog.getExistingDirectory(self._ui, self.fc_dict["open_m"][self.fc_lg], "", options = options)
            if presentDir:
                working_dir = os.getcwd()
                self._current_file_list.clear()
                self._file_id_list.clear()
                file_list = os.listdir(presentDir)
                file_final_list = []
                for file in file_list:
                    if file.endswith('.txt'):
                        file_version = re.search(r's0|t\d+|EN|ZH', file)
                        file_id = re.sub(r's0|t\d+|EN|ZH', "", file)
                        file_id = re.sub(r'\.txt', "", file_id)
                        if file_version:
                            file_final_list.append(file)
                            self._file_id_list.append(file_id)
                            file_path = os.path.join(presentDir, file)
                            file_path = file_path.replace('\\', "/")
                            self._current_file_list.append(file_path)
                        else:
                            pass
                    else:
                        pass
                if file_final_list:
                    self._ui._file_openBox.setText(",".join(file_final_list))
                    self.open_multi_file(self._current_file_list, self._file_id_list)
                else:
                    pass
            else:
                self._current_file_list.clear()
                self._file_id_list.clear()
                self._ui._set_status_text(self.fc_dict["warning_naming_error"][self.fc_lg])
        else:
            pass

    def sg_file_evaluator(self, para_list):
        lang_status_dict = {self.fc_dict['u_d'][self.fc_lg]: 0, self.fc_dict['l_r'][self.fc_lg]: 1, self.fc_dict["bi-sep"][self.fc_lg]: 2}
        if para_list == []:
            self._ui._set_status_text(self.fc_dict["warning_read_fail_unknown"][self.fc_lg])
        else:
            tab_test = para_list[0].split('\t')
            for text in tab_test:
                tg = self._lg_detect.detect_lang(text)
                if tg != 'num':
                    self.sl = tg
                    break
            if self.sl == 'zh':
                self.tl = 'en'
            else:
                self.sl = 'en'
                self.tl = 'zh'
            if len(tab_test) == 1:
                col_max = 1
                marker_id_status = 0
                self._ui._file_num_mark_box.setChecked(False)
                self._opt_dict['marker_id'] = 0
                marker_chapter = 0
                self._ui._file_chapt_num_box.setChecked(False)
                self._opt_dict['marker_chapt'] = 0
                self._ui._file_tab_mark_box.setChecked(False)
                self._opt_dict['marker_tab'] = 0
                lang_status, lang_gap, self.sl, self.tl = self._lg_detect.detect_lang_swap(self.fc_dict, self.fc_lg, para_list)
                self._ui._file_portion_box.setValue(lang_gap)
                self._opt_dict['lang_cols'] = lang_gap
                pos_index = lang_status_dict[lang_status]
                self._ui._file_pos_box.setCurrentIndex(pos_index)
                self._opt_dict['lang_pos'] = lang_status
                if lang_status == self.fc_dict['u_d'][self.fc_lg]:
                    self._opt_dict['lang_rows'] = int(self._ui._file_portion_box.value()) + 1
                    self._opt_dict['lang_cols'] = 1
                elif lang_status == self.fc_dict['l_r'][self.fc_lg]:
                    self._opt_dict['lang_rows'] = 1
                    self._opt_dict['lang_cols'] = int(self._ui._file_portion_box.value())
                elif lang_status == self.fc_dict["bi-sep"][self.fc_lg]:
                    self._opt_dict['lang_rows'] = 1
                    self._opt_dict['lang_cols'] = 0
                else:
                    pos_index = lang_status_dict[lang_status]
                    self._ui._file_pos_box.setCurrentIndex(pos_index)
            elif len(tab_test) >= 2:
                lang_status = self.fc_dict["l_r"][self.fc_lg]
                col_max = len(tab_test)
                self._ui._file_tab_mark_box.setChecked(True)
                self._opt_dict['marker_tab'] = 1
                tab_list = []
                if '' not in tab_test:
                    for item in para_list:
                        item_list = item.split('\t')
                        tab_list.append(item_list)
                else:
                    for item in para_list:
                        item_list = item.split('\t')
                        if '' in item_list:
                            pass
                        else:
                            tab_list.append(item_list)
                colum_max = len(tab_list[0])
                lang_seq_dict = {}
                lang_seq_list = []
                for i in range(colum_max):
                    target_lang = self._lg_detect.detect_lang(tab_list[0][i])
                    lang_seq_dict[target_lang] = i
                    lang_seq_list.append(target_lang)
                count_col_num = lang_seq_list.count('num')
                count_col_sl = lang_seq_list.count(self.sl)
                count_col_tl = lang_seq_list.count(self.tl)
                col_names = [x for x in lang_seq_dict.keys()]
                if self.sl in lang_seq_dict.keys() and self.tl in lang_seq_dict.keys():
                    lang_status = self.fc_dict['l_r'][self.fc_lg]
                    if lang_seq_dict[self.sl] < lang_seq_dict[self.tl]:
                        lang_status = self.fc_dict['l_r'][self.fc_lg]
                        if count_col_sl == 1:
                            lang_gap = count_col_tl
                            self._ui._file_portion_box.setValue(lang_gap)
                            self._opt_dict['lang_cols'] = lang_gap
                            self._opt_dict['lang_rows'] = 1
                            title = self.fc_dict["title_no"][self.fc_lg]
                        else:
                            lang_gap = int(count_col_tl / 2)
                            self._ui._file_portion_box.setValue(lang_gap)
                            self._opt_dict['lang_cols'] = lang_gap
                            self._opt_dict['lang_rows'] = 1
                            title = self.fc_dict["title_yes"][self.fc_lg]
                    else:
                        if count_col_tl == 1:
                            lang_gap = count_col_sl
                            self._ui._file_portion_box.setValue(lang_gap)
                            self._opt_dict['lang_cols'] = lang_gap
                            self._opt_dict['lang_rows'] = 1
                            title = self.fc_dict["title_no"][self.fc_lg]
                        else:
                            lang_gap = int(count_col_sl / 2)
                            self._ui._file_portion_box.setValue(lang_gap)
                            self._opt_dict['lang_cols'] = lang_gap
                            self._opt_dict['lang_rows'] = 1
                            title = self.fc_dict["title_yes"][self.fc_lg]
                # 否则为上下结构
                elif self.sl in lang_seq_dict.keys():
                    lang_status = self.fc_dict["u_d"][self.fc_lg]
                    if count_col_sl == 1:
                        title = self.fc_dict["title_no"][self.fc_lg]
                    else:
                        title = self.fc_dict["title_yes"][self.fc_lg]
                    col_sent_list = []
                    col_id = lang_seq_dict[self.sl]
                    for item in tab_list:
                        col_sent_list.append(item[col_id])
                    lang_status, lang_gap, self.sl, self.tl = self._lg_detect.detect_lang_swap(self.fc_dict, self.fc_lg, col_sent_list)
                    self._ui._file_portion_box.setValue(lang_gap)
                    #self._opt_dict['lang_cols'] = lang_gap
                    self._opt_dict['lang_rows'] = int(self._ui._file_portion_box.value()) + 1
                    self._opt_dict['lang_cols'] = 1
                elif self.tl in lang_seq_dict.keys():
                    if count_col_sl == 1:
                        title = self.fc_dict["title_no"][self.fc_lg]
                    else:
                        title = self.fc_dict["title_yes"][self.fc_lg]
                    col_sent_list = []
                    col_id = lang_seq_dict[self.tl]
                    for item in tab_list:
                        col_sent_list.append(item[col_id])
                    lang_status, lang_gap, self.sl, self.tl = self._lg_detect.detect_lang_swap(self.fc_dict, self.fc_lg, col_sent_list)
                    self._ui._file_portion_box.setValue(lang_gap)
                    #self._opt_dict['lang_cols'] = lang_gap
                    self._opt_dict['lang_rows'] = int(self._ui._file_portion_box.value()) + 1
                    self._opt_dict['lang_cols'] = 1
                else:
                    title = ''
                    lang_status = ''
                    lang_gap = ''
                pos_index = lang_status_dict[lang_status]
                self._ui._file_pos_box.setCurrentIndex(pos_index)
                self._opt_dict['lang_pos'] = lang_status
                self._opt_dict['lang_cols'] = lang_gap
                self._ui._file_portion_box.setValue(lang_gap)
                # 输出报告
                if count_col_num == 0:
                    num = self.fc_dict["line_num_no"][self.fc_lg]
                    marker_id_status = 0
                    self._ui._file_num_mark_box.setChecked(False)
                    self._opt_dict['marker_id'] = 0
                else:
                    num = self.fc_dict["line_num_yes"][self.fc_lg]
                    marker_id_status = 1
                    self._ui._file_num_mark_box.setChecked(True)
                    self._opt_dict['marker_id'] = 1
                if title == self.fc_dict["title_no"][self.fc_lg]:
                    marker_chapter = 0
                    self._ui._file_chapt_num_box.setChecked(False)
                    self._opt_dict['marker_chapt'] = 0
                else:
                    marker_chapter = 1
                    self._ui._file_chapt_num_box.setChecked(True)
                    self._opt_dict['marker_chapt'] = 1

            else:
                marker_id_status = -1
                marker_chapter = -1
                row_max = -1
                col_max = -1

            self._temp_list.clear()
            file_pos = self._opt_dict['lang_pos']
            row_max = self._opt_dict['lang_rows']
            if col_max >= 1:
                if file_pos == self.fc_dict['u_d'][self.fc_lg]:
                    self._list_preparer.prepare_return_bi_list(self._ui, self._current_sl_para_list, self._current_tl_para_list, self._temp_list,self._opt_dict, self.sl,self.tl, file_pos, row_max, col_max, para_list)
                elif file_pos == self.fc_dict['l_r'][self.fc_lg] and col_max >= 2:
                    if self._opt_dict['marker_chapt'] == 0:
                        col_max = self._opt_dict['marker_id'] + self._opt_dict['lang_cols'] + 1
                    elif self._opt_dict['marker_chapt'] == 1:
                        col_max = self._opt_dict['marker_id'] + (self._opt_dict['lang_cols'] + 1) * 2
                    else:
                        col_max = 0
                    self._list_preparer.prepare_tab_bi_list(self._ui, self._current_sl_para_list, self._current_tl_para_list, self._temp_list,self._opt_dict, self.sl, self.tl, file_pos, row_max, col_max, para_list)
                elif file_pos == self.fc_dict["bi-sep"][self.fc_lg]:
                    pass
                else:
                    self._ui._set_status_text(self.fc_dict["warning_read_fail_unknown"][self.fc_lg])
            else:
                pass
            return self.sl,self.tl,self._current_sl_para_list, self._current_tl_para_list
    # opt_organizer_group
    def sep_files_evaluator(self, para_list):
        if para_list == []:
            self._ui._set_status_text(self.fc_dict["warning_read_fail_unknown"][self.fc_lg])
        else:
            lang_status = 2
            lang_gap = 0
            self._ui._file_portion_box.setValue(lang_gap)
            self._ui._file_pos_box.setCurrentIndex(lang_status)
            self._opt_dict['lang_rows'] = 1
            self._opt_dict['lang_cols'] = 0
            tab_test = para_list[0].split('\t')
            for text in tab_test:
                tg = self._lg_detect.detect_lang(text)
                if tg != 'num':
                    self.sl = tg
                    break
            if self.sl == 'zh':
                self.tl = 'en'
            else:
                self.sl = 'en'
                self.tl = 'zh'
            if len(tab_test) == 1:
                self._opt_dict['marker_tab'] = 0
                self._ui._file_tab_mark_box.setChecked(False)
                self._opt_dict['marker_id'] = 0
                self._ui._file_num_mark_box.setChecked(False)
                self._opt_dict['marker_chapt'] = 0
                self._ui._file_chapt_num_box.setChecked(False)
            elif len(tab_test) == 2:
                self._opt_dict['marker_tab'] = 1
                self._ui._file_tab_mark_box.setChecked(True)
                if tab_test[0].isdigit() == True:
                    self._opt_dict['marker_id'] = 1
                    self._ui._file_num_mark_box.setChecked(True)
                    self._opt_dict['marker_chapt'] = 0
                    self._ui._file_chapt_num_box.setChecked(False)
                else:
                    self._opt_dict['marker_id'] = 0
                    self._ui._file_num_mark_box.setChecked(False)
                    self._opt_dict['marker_chapt'] = 1
                    self._ui._file_chapt_num_box.setChecked(True)
            elif len(tab_test) == 3:
                self._opt_dict['marker_tab'] = 1
                self._ui._file_tab_mark_box.setChecked(True)
                self._opt_dict['marker_id'] = 1
                self._ui._file_num_mark_box.setChecked(True)
                self._opt_dict['marker_chapt'] = 1
                self._ui._file_chapt_num_box.setChecked(True)
            else:
                pass
            self._temp_list.clear()
            return self.sl,self.tl

    # opt_organizer_group
    def opt_checker(self):
        file_dict = {}
        file_dict['file_path'] = self._ui._file_openBox.text()
        file_dict['file_type'] = self._ui._file_type_box.currentText()
        file_dict['file_bind'] = self._ui._file_bind_box.currentText()
        file_dict['file_num'] = int(self._ui._file_num_box.value())
        file_dict['lang_pos'] = self._ui._file_pos_box.currentText()
        if file_dict['lang_pos'] == self.fc_dict['u_d'][self.fc_lg]:
            file_dict['lang_rows'] = int(self._ui._file_portion_box.value()) + 1
            file_dict['lang_cols'] = 1
        elif file_dict['lang_pos'] == self.fc_dict['l_r'][self.fc_lg]:
            file_dict['lang_cols'] = int(self._ui._file_portion_box.value())
            file_dict['lang_rows'] = 1
        else:
            file_dict['lang_cols'] = -1
            file_dict['lang_rows'] = -1
        if self._ui._file_num_mark_box.isChecked():
            file_dict['marker_id'] = 1
        else:
            file_dict['marker_id'] = 0
        if self._ui._file_chapt_num_box.isChecked():
            file_dict['marker_chapt'] = 1
        else:
            file_dict['marker_chapt'] = 0
        if self._ui._file_tab_mark_box.isChecked():
            file_dict['marker_tab'] = 1
        else:
            file_dict['marker_tab'] = 0
        if self._ui._file_cuc_mark_box.isChecked():
            file_dict['marker_seg'] = 1
        else:
            file_dict['marker_seg'] = 0
        if self._ui._file_table_mark_box.isChecked():
            file_dict['marker_table'] = 1
        else:
            file_dict['marker_table'] = 0

        return file_dict

    # auto_filler_group
    def json_sl_submitter(self):
        self._current_sl_chapter_num_list.clear()
        alert_msg = []
        if self._ui._file_openBox.text() == "":
            alert_msg.append(self.fc_dict["content"][self.fc_lg])
        if self._ui._ss_book_titleBox.text() == "":
            alert_msg.append(self.fc_dict["title_book"][self.fc_lg])
        if self._ui._ss_book_authorBox.text() == "":
            alert_msg.append(self.fc_dict["ar_id"][self.fc_lg])
        if self._ui._ss_book_genreBox.text() == "":
            alert_msg.append(self.fc_dict["issue_genre"][self.fc_lg])
        if alert_msg:
            alert_msg = " ".join(alert_msg)
            alert_detail = self.fc_dict["lg_sl"][self.fc_lg]  + " " + alert_msg + " " + self.fc_dict["warning_blank_error"][self.fc_lg]
            self._ui._set_status_text(alert_detail.replace("：",''))
        else:
            title = self._ui._ss_book_titleBox.text()
            self._current_dict_key = title.strip()
            self._current_dict_key = re.sub(r'\s+', '', self._current_dict_key)
            author = self._ui._ss_book_authorBox.text()
            translator = ''
            language = self._ui._ss_book_languageBox.text()
            self.sl = language
            if self.sl == 'zh':
                self.tl == 'en'
            date = self._ui._ss_book_dateBox.text()
            genre = self._ui._ss_book_genreBox.text()
            version = self._ui._ss_book_versionBox.text()
            self._current_sl_dict_version = version
            current_sl_text = self._ui._ss_book_contentsBox.toPlainText()
            current_sl_text_list = current_sl_text.split('\n')
            line_sample = current_sl_text_list[0]
            line_part = line_sample.split('\t')
            col_count = len(line_part)
            if col_count == 1:
                temp_num_list = []
                temp_sent_list = current_sl_text_list
                temp_chapter_list = []
            elif col_count == 2:
                if line_sample[0].isdigit() == True:
                    temp_num_list = [sent.split('\t')[0] for sent in current_sl_text_list]
                    temp_sent_list = [sent.split('\t')[1] for sent in current_sl_text_list]
                    temp_chapter_list = []
                else:
                    temp_num_list = []
                    temp_sent_list = [sent.split('\t')[0] for sent in current_sl_text_list]
                    temp_chapter_list = [sent.split('\t')[1] for sent in current_sl_text_list]
            else:
                temp_num_list = [sent.split('\t')[0] for sent in current_sl_text_list]
                temp_sent_list = [sent.split('\t')[1] for sent in current_sl_text_list]
                temp_chapter_list = [sent.split('\t')[2] for sent in current_sl_text_list]

            if temp_num_list == [] and temp_chapter_list == []:
                new_sent_list, new_chapter_list, sl_chpt_num_list = self._chapter_maker.add_chapter(self.sl,self.tl,temp_sent_list)
                sl_contents = [str(num) + "\t" + para + "\t" + chapter for num, (para, chapter) in
                               enumerate(zip(new_sent_list, new_chapter_list), start = 1)]
                sl_contents_shown = [str(num) + " ¦ " + para + ' ¦ ' + chapter for num, (para, chapter) in
                                     enumerate(zip(new_sent_list, new_chapter_list), start = 1)]
                self._current_sl_chapter_num_list = sl_chpt_num_list
            elif temp_num_list == []:
                sl_contents = [str(num) + '\t' + sent + '\t' + chapter for num, (sent, chapter) in
                               enumerate(zip(temp_sent_list, temp_chapter_list), start = 1)]
                sl_contents_shown = [str(num) + " ¦ " + sent + ' ¦ ' + chapter for num, (sent, chapter) in
                                     enumerate(zip(temp_sent_list, temp_chapter_list), start = 1)]
                self._current_sl_chapter_num_list = []
            elif temp_chapter_list == []:
                new_sent_list, new_chapter_list, sl_chpt_num_list = self._chapter_maker.add_chapter(self.sl,self.tl,temp_sent_list)
                sl_contents = [str(num) + '\t' + sent + '\t' + chapter for num, sent, chapter in
                               zip(temp_num_list, new_sent_list, new_chapter_list)]
                sl_contents_shown = [str(num) + " ¦ " + sent + ' ¦ ' + chapter for num, sent, chapter in
                                     zip(temp_num_list, new_sent_list, new_chapter_list)]
                self._current_sl_chapter_num_list = sl_chpt_num_list
            else:
                sl_contents = [str(num) + '\t' + sent + '\t' + chapter for num, sent, chapter in
                               zip(temp_num_list, temp_sent_list, temp_chapter_list)]
                sl_contents_shown = [str(num) + " ¦ " + sent + ' ¦ ' + chapter for num, sent, chapter in
                                     zip(temp_num_list, temp_sent_list, temp_chapter_list)]
                self._current_sl_chapter_num_list = []

            self._ui._ss_book_contentsBox.clear()
            self._ui._ss_book_contentsBox.setText("\n".join(sl_contents_shown))
            self._sl_bookDict.clear()
            self._sl_bookDict['title'] = title
            self._sl_bookDict['author'] = author
            self._sl_bookDict['translator'] = translator
            self._sl_bookDict['language'] = language
            self._sl_bookDict['date'] = date
            self._sl_bookDict['genre'] = genre
            self._sl_bookDict['content'] = defaultdict(str)
            for line in sl_contents:
                line = line.split('\t')
                num = line[0]
                para = line[1]
                chapt = line[2]
                self._sl_bookDict['content'][num] = para + '\t' + chapt

            self._ui._ss_book_uploadButton.setEnabled(False)
            self._ui._tt_book_uploadButton.setEnabled(True)
            self._ui._ss_book_redoButton.setEnabled(True)
            self._ui._promptBox.setText(self._ui._prompt_3)

        return self._current_dict_key, self._current_sl_dict_version, self._current_sl_chapter_num_list

    # auto_filler_group
    def json_tl_submitter(self):
        alert_msg = []
        if self._ui._file_openBox.text() == "":
            alert_msg.append(self.fc_dict["content"][self.fc_lg])
        if self._ui._tt_book_titleBox.text() == "":
            alert_msg.append(self.fc_dict["title_book"][self.fc_lg])
        if self._ui._tt_book_translatorBox.text() == "":
            alert_msg.append(self.fc_dict["tr_id"][self.fc_lg])
        if self._ui._tt_book_genreBox.text() == "":
            alert_msg.append(self.fc_dict["issue_genre"][self.fc_lg])
        if alert_msg:
            alert_msg = " ".join(alert_msg)
            alert_detail = self.fc_dict["lg_tl"][self.fc_lg] + " " + alert_msg + self.fc_dict["warning_blank_error"][self.fc_lg]
            self._ui._set_status_text(alert_detail.replace("：",''))
        else:
            title = self._ui._tt_book_titleBox.text()
            author = self._ui._tt_book_authorBox.text()
            translator = self._ui._tt_book_translatorBox.text()
            language = self._ui._tt_book_languageBox.text()
            date = self._ui._tt_book_dateBox.text()
            genre = self._ui._tt_book_genreBox.text()
            version = self._ui._tt_book_versionBox.text()
            self._current_tl_dict_version = version
            tl_vn = int(version.replace("t", ""))
            version_count = len(self._current_tl_para_list)
            current_tl_text_list = self._current_tl_para_list[tl_vn - 1]
            if tl_vn < version_count:
                self._ui._tt_book_nextButton.setEnabled(True)
            else:
                self._ui._tt_book_nextButton.setEnabled(False)
                self._ui._tt_book_uploadButton.setEnabled(False)

            line_sample = current_tl_text_list[0]
            # 列表内容如果是元组：
            if isinstance(line_sample, tuple) == True:
                # 统计元组元素个数用len方法
                col_count = len(line_sample)
                if col_count == 1:
                    temp_num_list = []
                    temp_sent_list = [x for x in current_tl_text_list if x != '']
                    temp_chapter_list = []
                elif col_count == 2:
                    if line_sample[0].isdigit() == True:
                        temp_num_list = [sent[0] for sent in current_tl_text_list if sent[1] != '']
                        temp_sent_list = [sent[1] for sent in current_tl_text_list if sent[1] != '']
                        temp_chapter_list = []
                    else:
                        temp_num_list = []
                        temp_sent_list = [sent[0] for sent in current_tl_text_list if sent[0] != '']
                        temp_chapter_list = [sent[1] for sent in current_tl_text_list if sent[1] != '']
                else:
                    temp_num_list = [sent[0] for sent in current_tl_text_list if sent[1] != '']
                    temp_sent_list = [sent[1] for sent in current_tl_text_list if sent[1] != '']
                    temp_chapter_list = [sent[2] for sent in current_tl_text_list if sent[2] != '']
            # 否则为字符串：
            else:
                col_count = len(line_sample.split('\t'))
                if col_count == 1:
                    temp_num_list = []
                    temp_sent_list = [x for x in current_tl_text_list if x != '']
                    temp_chapter_list = []
                elif col_count == 2:
                    if line_sample[0].isdigit() == True:
                        temp_num_list = [sent.split('\t')[0] for sent in current_tl_text_list if
                                         sent.split('\t')[1] != '']
                        temp_sent_list = [sent.split('\t')[1] for sent in current_tl_text_list if
                                          sent.split('\t')[1] != '']
                        temp_chapter_list = []
                    else:
                        temp_num_list = []
                        temp_sent_list = [sent.split('\t')[0] for sent in current_tl_text_list if
                                          sent.split('\t')[0] != '']
                        temp_chapter_list = [sent.split('\t')[1] for sent in current_tl_text_list if
                                             sent.split('\t')[1] != '']
                else:
                    temp_num_list = [sent.split('\t')[0] for sent in current_tl_text_list if sent.split('\t')[1] != '']
                    temp_sent_list = [sent.split('\t')[1] for sent in current_tl_text_list if sent.split('\t')[1] != '']
                    temp_chapter_list = [sent.split('\t')[2] for sent in current_tl_text_list if
                                         sent.split('\t')[2] != '']

            if temp_num_list == [] and temp_chapter_list == []:
                temp_num_list = [num for num, sent in enumerate(temp_sent_list, start = 1)]
                temp_sent_chapter_list = self._chapter_maker.swap_chapter(zip(temp_num_list, temp_sent_list),
                                                           self._current_sl_chapter_num_list)
                tl_contents = [str(num) + "\t" + para + "\t" + chapter for num, (para, chapter) in
                               zip(temp_num_list, temp_sent_chapter_list)]
                tl_contents_shown = [str(num) + " ¦ " + para + ' ¦ ' + chapter for num, (para, chapter) in
                                     zip(temp_num_list, temp_sent_chapter_list)]
            elif temp_num_list == []:
                tl_contents = [str(num) + '\t' + sent + '\t' + chapter for num, (sent, chapter) in
                               enumerate(zip(temp_sent_list, temp_chapter_list), start = 1)]
                tl_contents_shown = [str(num) + " ¦ " + sent + ' ¦ ' + chapter for num, (sent, chapter) in
                                     enumerate(zip(temp_sent_list, temp_chapter_list), start = 1)]
            elif temp_chapter_list == []:

                temp_sent_chapter_list = self._chapter_maker.swap_chapter(zip(temp_num_list, temp_sent_list),
                                                           self._current_sl_chapter_num_list)
                tl_contents = [str(num) + '\t' + sent + '\t' + chapter for num, (sent, chapter) in
                               zip(temp_num_list, temp_sent_chapter_list)]
                tl_contents_shown = [str(num) + " ¦ " + sent + ' ¦ ' + chapter for num, (sent, chapter) in
                                     zip(temp_num_list, temp_sent_chapter_list)]
            else:
                tl_contents = [str(num) + '\t' + sent + '\t' + chapter for num, sent, chapter in
                               zip(temp_num_list, temp_sent_list, temp_chapter_list)]
                tl_contents_shown = [str(num) + " ¦ " + sent + ' ¦ ' + chapter for num, sent, chapter in
                                     zip(temp_num_list, temp_sent_list, temp_chapter_list)]

            self._ui._tt_book_contentsBox.clear()
            self._ui._tt_book_contentsBox.setText("\n".join(tl_contents_shown))
            tl_bookDict = {}
            tl_bookDict['title'] = title
            tl_bookDict['author'] = author
            tl_bookDict['translator'] = translator
            tl_bookDict['language'] = language
            tl_bookDict['date'] = date
            tl_bookDict['genre'] = genre
            tl_bookDict['content'] = defaultdict(str)
            for line in tl_contents:
                line = line.split('\t')
                num = line[0]
                para = line[1]
                chapt = line[2]
                tl_bookDict['content'][num] = para + '\t' + chapt
            self._tl_bookDict_list.append(tl_bookDict)
            self._ui._tt_book_redoButton.setEnabled(True)
            self._ui._tt_book_uploadButton.setEnabled(False)
            self._ui._prompt_4 = self.fc_dict["corp_tl"][self.fc_lg] + f"{self._current_tl_dict_version}" + \
                                     self.fc_dict["pmt_4a_start"][self.fc_lg] + "\n" + self.fc_dict["pmt_4a_a"][
                                         self.fc_lg]
            self._ui._prompt_4b = self.fc_dict["corp_tl"][self.fc_lg] + f"{self._current_tl_dict_version}" + \
                                      self.fc_dict["pmt_4b_start"][self.fc_lg] + "\n" + self.fc_dict["pmt_4b_a"][
                                          self.fc_lg]
            if self._ui._tt_book_nextButton.isEnabled() == False:
                self._ui._promptBox.setText(self._ui._prompt_4)
            else:
                self._ui._promptBox.setText(self._ui._prompt_4b)

        return self._current_tl_dict_version

    def json_tl_vn_loader(self):
        title = self._ui._tt_book_titleBox.text()
        author = self._ui._tt_book_authorBox.text()
        translator = self._ui._tt_book_translatorBox.text()
        language = self._ui._tt_book_languageBox.text()
        date = self._ui._tt_book_dateBox.text()
        genre = self._ui._tt_book_genreBox.text()
        version = self._ui._tt_book_versionBox.text()
        self._current_tl_dict_version = version
        book_id = title
        chapter = title.strip()
        tl_vn = int(version.replace("t", ""))
        version_count = len(self._current_tl_para_list)
        current_tl_text_list = self._current_tl_para_list[tl_vn - 1]
        if tl_vn < version_count:
            self._ui._tt_book_nextButton.setEnabled(True)
        else:
            self._ui._tt_book_nextButton.setEnabled(False)
            self._ui._tt_book_uploadButton.setEnabled(False)

        self._ui._tt_book_contentsBox.clear()
        self._ui._tt_book_contentsBox.setText("\n".join(current_tl_text_list))

        return self._current_tl_dict_version, current_tl_text_list

    def option_reset(self):
        self.sl = 'en'
        self.tl = 'zh'
        self._ui._file_openBox.clear()
        self._ui._file_type_box.setCurrentIndex(0)
        if self._current_file:
            self._ui._file_bind_box.setCurrentIndex(0)
        else:
            self._ui._file_bind_box.setCurrentIndex(1)
        self._ui._file_num_box.setValue(1)
        self._ui._file_portion_box.setValue(1)
        self._ui._file_pos_box.setCurrentIndex(0)
        self._ui._file_num_mark_box.setChecked(False)
        self._ui._file_chapt_num_box.setChecked(False)
        self._ui._file_tab_mark_box.setChecked(False)
        self._ui._file_cuc_mark_box.setChecked(False)
        self._ui._file_table_mark_box.setChecked(False)

    def form_reset(self):
        self._ui._ss_book_titleBox.clear()
        self._ui._tt_book_titleBox.clear()
        self._ui._ss_book_authorBox.clear()
        self._ui._tt_book_authorBox.clear()
        self._ui._tt_book_translatorBox.clear()
        self._ui._ss_book_genreBox.clear()
        self._ui._tt_book_genreBox.clear()
        self._ui._ss_book_dateBox.clear()
        self._ui._tt_book_dateBox.clear()
        self._ui._tt_book_versionBox.setText('t1')
        self._ui._ss_book_contentsBox.clear()
        self._ui._tt_book_contentsBox.clear()
        self._ui._ss_book_uploadButton.setEnabled(True)
        self._ui._tt_book_uploadButton.setEnabled(False)
        self._ui._ss_book_redoButton.setEnabled(False)
        self._ui._tt_book_redoButton.setEnabled(False)
        self._current_sl_para_list.clear()
        self._current_tl_para_list.clear()
        self._sl_bookDict = {}
        self._tl_bookDict_list.clear()
        self._temp_list.clear()

        return self._sl_bookDict    
     
    # auto_refiller_group
    def json_sl_refiller(self):
        self._ui._ss_book_uploadButton.setEnabled(True)
        self._sl_bookDict.clear()
        self._ui._ss_book_titleBox.clear()
        self._ui._ss_book_authorBox.clear()
        self._ui._ss_book_genreBox.clear()
        self._ui._ss_book_dateBox.clear()
        self._ui._ss_book_contentsBox.clear()
        self._current_sl_dict_version = ""
        self._current_dict_key = ""
        if self._current_sl_para_list:
            self._ui._ss_book_contentsBox.setText("\n".join(self._current_sl_para_list))
        else:
            pass

        return self._current_dict_key, self._current_sl_dict_version

    # auto_refiller_group
    def json_tl_refiller(self):
        self._ui._tt_book_uploadButton.setEnabled(True)
        self._tl_bookDict_list.clear()
        self._ui._tt_book_titleBox.clear()
        self._ui._tt_book_authorBox.clear()
        self._ui._tt_book_translatorBox.clear()
        self._ui._tt_book_genreBox.clear()
        self._ui._tt_book_versionBox.clear()
        self._ui._tt_book_versionBox.setText('t1')
        self._ui._tt_book_dateBox.clear()
        self._ui._tt_book_contentsBox.clear()
        if self._current_tl_para_list:
            self._ui._tt_book_contentsBox.setText("\n".join(self._current_tl_para_list[0]))
        else:
            pass
        self._current_tl_dict_version = 't1'
        return self._current_tl_dict_version

    # version_shift_group
    def load_next_version (self):
        version_count = len(self._current_tl_para_list)
        version_num = int(self._current_tl_dict_version.replace("t", ""))
        if version_num <= version_count:
            next_num = version_num + 1
        else:
            next_num = ""
        if next_num:
            self._ui._tt_book_versionBox.setText(f't{next_num}')
            current_tl_vn, current_tl_text = self.json_tl_vn_loader()
            temp_text = current_tl_text
            tab_test = temp_text[0].split('\t')
            tab_units = len(tab_test)
            if tab_units == 1:
                temp_lines = temp_text[:3]
            elif tab_units >= 2:
                temp_lines = [item.split('\t')[1] for item in temp_text[:3]]
            else:
                temp_lines = []
            self.sl = self._ui._ss_book_languageBox.text()
            self.tl = self._ui._tt_book_languageBox.text()
            if self.tl == "zh":
                tmp_tl_title, tmp_tl_author, tmp_tl_translator, tmp_tl_date = self._info_collector.info_collector_zh(self.sl,self.tl,temp_lines)
            else:
                tmp_tl_title, tmp_tl_author, tmp_tl_translator, tmp_tl_date = self._info_collector.info_collector_en(
                    self.sl, self.tl, temp_lines)
            try:
                self._ui._tt_book_titleBox.setText(tmp_tl_title)
                self._ui._tt_book_translatorBox.setText(tmp_tl_translator)
                if tmp_tl_author:
                    self._ui._tt_book_authorBox.setText(tmp_tl_author)
                else: pass
                self._ui._tt_book_dateBox.setText(tmp_tl_date)
                self._ui._prompt_4a = self.fc_dict["pmt_4c_start"][self.fc_lg]+f"t{next_num}， "+self.fc_dict["pmt_4c_a"][self.fc_lg]
                self._ui._promptBox.setText(self._ui._prompt_4a)
            except:
                self._ui._tt_book_translatorBox.clear()
                self._ui._tt_book_dateBox.clear()
            finally:
                self._ui._tt_book_uploadButton.setEnabled(True)
                self._ui._tt_book_nextButton.setEnabled(False)
        else:
            pass

    # data_save_group
    def write_to_json(self):
        if self._sl_bookDict and self._tl_bookDict_list:
            self.sl = self._sl_bookDict['language']
            self.tl = self._tl_bookDict_list[0]['language']
            mydict = {}
            book_id = self._current_dict_key
            if book_id:
                sl_version = self._current_sl_dict_version
                mydict[book_id] = {}
                mydict[book_id][self.sl] = {}
                mydict[book_id][self.tl] = {}
                mydict[book_id][self.sl][sl_version] = self._sl_bookDict
                for i, tl_version_dict in enumerate(self._tl_bookDict_list, start = 1):
                    tl_version = "t" + str(i)
                    mydict[book_id][self.tl][tl_version] = tl_version_dict
                json_to_save = os.path.join(self._outPutDir, book_id + ".json")
                if mydict:
                    with open(json_to_save, mode = 'w', encoding = "utf-8-sig") as f:
                        json.dump(mydict, f)
                        self._ui._promptBox.setText(self._ui._prompt_5)
                else:
                    pass
            else:
                pass
        else:
            self._ui._set_status_text(self.fc_dict["warning_not_ready"][self.fc_lg])
